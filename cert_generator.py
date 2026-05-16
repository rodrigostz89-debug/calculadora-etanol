import io
import re
from datetime import datetime
import PyPDF2
from docx import Document

def extrair_dados_pdf(file_obj):
    texto = ""
    try:
        reader = PyPDF2.PdfReader(file_obj)
        for page in reader.pages:
            if page.extract_text():
                texto += page.extract_text() + "\n"
    except Exception as e:
        # Se falhar como PDF, tenta ler como texto/HTML (muito comum quando o navegador salva a página)
        try:
            file_obj.seek(0)
            texto = file_obj.read().decode('utf-8', errors='ignore')
            # Remove tags HTML simples para limpar o texto
            texto = re.sub(r'<[^>]+>', ' ', texto)
        except Exception as e2:
            print(f"Erro ao ler arquivo: {e2}")
            return "", ""

    # Extrair Quantidade (mais flexível)
    quant_match = re.search(r"Quantidade\s*[:\s]*([\d\.,]+)", texto, re.IGNORECASE)
    quantidade = quant_match.group(1).strip() if quant_match else ""
    if quantidade.endswith(",000"):
        quantidade = quantidade[:-4]

    # Extrair Placa
    placas_full = ""
    placa_match = re.search(r"Placas\s*[:\s]*([A-Z0-9\s-]+)", texto, re.IGNORECASE)
    if placa_match:
        placas_full = placa_match.group(1).strip()
    
    # Isolar a primeira placa e formatar com traço (Ex: TKL-0B33)
    primeira_placa = ""
    if placas_full:
        # Procura por um padrão de 3 letras e 4 números/letras (Mercosul ou Antiga)
        match_placa = re.search(r'([A-Z]{3})[- \.]?([0-9][A-Z0-9][0-9]{2})', placas_full, re.IGNORECASE)
        if match_placa:
            primeira_placa = f"{match_placa.group(1).upper()}-{match_placa.group(2).upper()}"
        else:
            # Fallback se não encontrar o padrão exato
            primeira_placa = placas_full.split()[0] if placas_full else ""
            
    return primeira_placa, quantidade

def preencher_certificado(template_file, numero_cert, placa, quantidade, data):
    doc = Document(template_file)
    
    substituicoes = {
        "{{NUMERO}}": numero_cert,
        "{{DATA}}": data,
        "{{PLACA}}": placa,
        "{{QUANTIDADE}}": quantidade
    }
    
    # Substituir em parágrafos comuns
    for p in doc.paragraphs:
        for tag, valor in substituicoes.items():
            if tag in p.text:
                # Substituição preservando o run o máximo possível ou resetando
                for run in p.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, str(valor))
                # Fallback caso a tag esteja dividida em múltiplos runs
                if tag in p.text:
                    p.text = p.text.replace(tag, str(valor))
                
    # Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for tag, valor in substituicoes.items():
                        if tag in p.text:
                            for run in p.runs:
                                if tag in run.text:
                                    run.text = run.text.replace(tag, str(valor))
                            if tag in p.text:
                                p.text = p.text.replace(tag, str(valor))
                            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io
